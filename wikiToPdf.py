import sys
import wikipedia
import re
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT 
from docx.shared import Cm,Pt
from docx.enum.style import WD_STYLE_TYPE
from docx2pdf import convert
import os

class Wikipedia():
    page_data = list()
    title = str()
    def __init__(self,search_input:str=None,lang:str=None) -> None:
        if not lang: # First set language
            lang = input("Set language: ")
        self.set_language(lang)
        
        print()
        print("Results")
        result_counter = 1
        results = self.search(search_input)
        for result in results:
            print(str(result_counter) + ".",result)
            result_counter += 1
        try:
            selection = int(input("Enter result number:"))
        except:
            print("Invalid input.. quiting")
            quit()
        
        page_content = self.get_content(results[selection-1])

        
    def set_language(self,lang:str):
        wikipedia.set_lang(lang)

    def search(self,search_input:str):
        results = wikipedia.search(search_input)
        return results

    def get_content(self,page):  
        self.title = wikipedia.page(page).title # Get title 
        summary = wikipedia.page(page).summary # Get first paragraph
        self.page_data.append(["t",self.title])
        #self.page_data.append(["p",summary])

        text = wikipedia.page(page).content # Get all other paragraphs
        text = text.split("\n\n")
        for value in text:
            heading = re.findall("== .* ==",value)
            if heading:
                heading = heading[0]
                heading = heading.replace("==","").rstrip().lstrip()
            parag = re.sub("== .* ==","",value)
            parag = parag.replace("\n==","").lstrip().rstrip()
            if parag: 
                self.page_data.append(["h",heading])
                self.page_data.append(["p",parag])
    
    def get_page_data(self):
        return self.page_data



class Document():
    doc = None
    def __init__(self,PATH:str):
        PATH = "output/" + PATH
        self.path = PATH
        self.doc = docx.Document()
        sections  = self.doc.sections
        self.__set_styles()
        for section in sections:
            section.top_margin = Cm(0.2)
            section.bottom_margin = Cm(0.2)
            section.left_margin = Cm(0.3)
            section.right_margin = Cm(0.3)
        
    def __set_styles(self):
        # Title style
        font_styles = self.doc.styles
        tit_style = font_styles.add_style('title', WD_STYLE_TYPE.CHARACTER)
        font_object = tit_style.font
        font_object.size = Pt(60)
        font_object.name = 'Times New Roman'

        #Heading style
        font_styles = self.doc.styles
        h_style = font_styles.add_style('heading', WD_STYLE_TYPE.CHARACTER)
        font_object = h_style.font
        font_object.size = Pt(30)
        font_object.name = 'Times New Roman'

        # Paragraph style
        font_styles = self.doc.styles
        p_style = font_styles.add_style('paragraph', WD_STYLE_TYPE.CHARACTER)
        font_object = p_style.font
        font_object.size = Pt(25)
        font_object.name = 'Times New Roman'
    
    def add_title(self,title_text:str):
        title = self.doc.add_paragraph("")
        title.add_run(title_text,style = 'title')
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def add_heading(self,heading_text:str):
        heading = self.doc.add_paragraph("")
        heading.add_run(heading_text,style = 'heading')
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def add_parag(self,parag_text:str):
        p = self.doc.add_paragraph("")
        p.add_run(parag_text,style = 'paragraph')
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    

    def save_document(self):
        if os.path.exists(self.path): # If file exists, overwrite it
            os.remove(self.path)
        self.doc.save(self.path)
        convert(self.path) # Convert file to pdf
        os.remove(self.path)


# python wikiToPdf.py [wiki page] [language]
wiki_page = str()
try:
    i = 1
    while True:
        wiki_page += sys.argv[i]
        wiki_page += " "
        i += 1
except:
    pass


wiki = Wikipedia(wiki_page)
page_data = wiki.page_data
page_name = wiki.title

doc = Document(page_name + ".docx")
for value in page_data:
    if value[0] == "p":
        doc.add_parag(value[1])
    elif value[0] == "h":
        doc.add_heading(value[1])
    elif value[0] == "t":
        doc.add_title(value[1])
doc.save_document()



